# PC06 OCR Engine - Tesseract Integration
# Nhẹ, nhanh, không cần GPU
# Designed for SERVER (auto-detect Tesseract)

import os

# Lazy imports
cv2 = None
fitz = None
Document = None
pytesseract = None

def _ensure_imports():
    global cv2, fitz, Document, pytesseract
    if cv2 is None:
        try:
            import cv2
        except: pass
    if fitz is None:
        try:
            import fitz
        except: pass
    if Document is None:
        try:
            from docx import Document
        except: pass
    if pytesseract is None:
        try:
            import pytesseract
        except: pass

# ==================== CẤU HÌNH TESSERACT ====================
# Tự động phát hiện - KHÔNG cần sửa thủ công!

def get_tesseract_path():
    """Tự động tìm đường dẫn Tesseract trên server"""
    import subprocess
    
    # Danh sách đường dẫn thông dụng theo OS
    common_paths = [
        # Windows
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        # Linux
        '/usr/bin/tesseract',
        '/usr/local/bin/tesseract',
        '/opt/bin/tesseract',
        # macOS  
        '/opt/homebrew/bin/tesseract',
        '/usr/local/bin/tesseract',
    ]
    
    # 1. Thử từng đường dẫn cụ thể
    for path in common_paths:
        if os.path.exists(path):
            print(f"[OCR] Found Tesseract at: {path}")
            return path
    
    # 2. Thử lệnh which (Linux/macOS)
    try:
        result = subprocess.run(['which', 'tesseract'], 
                         capture_output=True, text=True, timeout=5)
        if result.returncode == 0 and result.stdout.strip():
            path = result.stdout.strip()
            print(f"[OCR] Found Tesseract via which: {path}")
            return path
    except:
        pass
    
    # 3. Thử where (Windows)
    try:
        result = subprocess.run(['where', 'tesseract'],
                         capture_output=True, text=True, shell=True, timeout=5)
        if result.returncode == 0 and result.stdout.strip():
            path = result.stdout.strip().split('\n')[0].strip()
            print(f"[OCR] Found Tesseract via where: {path}")
            return path
    except:
        pass
    
    # 4. Để pytesseract tự tìm trong PATH
    print("[OCR] Will try to use Tesseract from system PATH")
    return None

# Ngôn ngữ OCR (vie = Tiếng Việt, eng = English)
TESSERACT_LANG = 'vie+eng'
TESSERACT_CONFIG = '--oem 3 --psm 6'

class PC06_Tesseract_OCR:
    def __init__(self):
        self.ocr_available = False
        self.status_message = ""
        
        _ensure_imports()
        
        # Kiểm tra các thư viện cơ bản
        has_cv2 = cv2 is not None
        has_docx = Document is not None
        has_fitz = fitz is not None
        has_pytesseract = pytesseract is not None
        
        if not has_pytesseract:
            self.status_message = "⚠️ Chưa cài pytesseract"
            print(f"[OCR] {self.status_message}")
            return
        
        if not has_docx:
            print("[OCR] ⚠️ Chưa cài python-docx")
        if not has_fitz:
            print("[OCR] ⚠️ Chưa cài pymupdf")
        
        # Tìm Tesseract
        tess_path = get_tesseract_path()
        
        if tess_path and os.path.exists(tess_path):
            pytesseract.pytesseract.tesseract_cmd = tess_path
            try:
                version = pytesseract.get_tesseract_version()
                self.ocr_available = True
                self.status_message = f"✅ Tesseract {version} sẵn sàng!"
                print(f"[OCR] {self.status_message}")
            except Exception as e:
                self.status_message = f"⚠️ Lỗi Tesseract: {e}"
                print(f"[OCR] {self.status_message}")
        else:
            # Thử chạy mà không cần đường dẫn cụ thể
            try:
                version = pytesseract.get_tesseract_version()
                self.ocr_available = True
                self.status_message = f"✅ Tesseract {version} (from PATH)"
                print(f"[OCR] {self.status_message}")
            except Exception as e:
                self.status_message = f"⚠️ Tesseract not found: {e}"
                print(f"[OCR] {self.status_message}")
    
    def preprocess_image(self, img_path):
        """Tiền xử lý ảnh để OCR tốt hơn"""
        _ensure_imports()
        if cv2 is None:
            return None
            
        try:
            img = cv2.imread(img_path)
            if img is None:
                return None
            
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            blurred = cv2.GaussianBlur(gray, (5, 5), 0)
            thresh = cv2.adaptiveThreshold(
                blurred, 255, 
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY, 11, 2
            )
            return thresh
        except Exception as e:
            print(f"[OCR] Preprocess error: {e}")
            return None
    
    def process_image_to_excel(self, img_path, output_path):
        """Chuyển ảnh có bảng sang Excel"""
        _ensure_imports()
        import pandas as pd
        
        if pytesseract is None:
            return False
        
        try:
            processed_img = self.preprocess_image(img_path)
            if processed_img is None:
                return False
            
            text = pytesseract.image_to_string(
                processed_img, 
                lang=TESSERACT_LANG,
                config=TESSERACT_CONFIG
            )
            
            if not text.strip():
                return False
            
            lines = text.strip().split('\n')
            data = []
            for line in lines:
                if line.strip():
                    parts = [p.strip() for p in line.split() if p.strip()]
                    if parts:
                        data.append(parts)
            
            if data:
                df = pd.DataFrame(data)
                if not df.empty:
                    df.to_excel(output_path, index=False, header=False)
                    return True
            
            return False
        except Exception as e:
            print(f"[OCR] Image to Excel error: {e}")
            return False
    
    def process_image_to_text(self, img_path):
        """Chuyển ảnh sang text"""
        _ensure_imports()
        
        if pytesseract is None:
            return "[Chưa cài Tesseract]"
        
        try:
            processed_img = self.preprocess_image(img_path)
            if processed_img is None:
                return "[Cannot read image]"
            
            text = pytesseract.image_to_string(
                processed_img,
                lang=TESSERACT_LANG,
                config=TESSERACT_CONFIG
            )
            
            return text.strip() if text.strip() else "[No text detected]"
        except Exception as e:
            return f"[Error: {str(e)}]"
    
    def process_pdf_to_word(self, pdf_path, output_word):
        """PDF sang Word"""
        _ensure_imports()
        if fitz is None or Document is None:
            return False
        
        try:
            doc = fitz.open(pdf_path)
            word_doc = Document()
            word_doc.add_heading('TRÍCH XUẤT TỪ PDF - PC06', 0)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text().strip()
                
                if text:
                    word_doc.add_paragraph(f"--- Trang {page_num+1} ---")
                    word_doc.add_paragraph(text)
                else:
                    word_doc.add_paragraph(f"--- Trang {page_num+1} (PDF scan) ---")
            
            word_doc.save(output_word)
            return True
        except Exception as e:
            print(f"[OCR] PDF to Word error: {e}")
            return False
    
    def process_pdf_to_excel(self, pdf_path, output_path):
        """PDF sang Excel"""
        _ensure_imports()
        if fitz is None:
            return False
        
        try:
            import pandas as pd
            doc = fitz.open(pdf_path)
            
            all_data = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text().strip()
                if text:
                    all_data.append({'Trang': page_num+1, 'Noi_dung': text[:500]})
            
            if all_data:
                df = pd.DataFrame(all_data)
                df.to_excel(output_path, index=False)
                return True
            return False
        except Exception as e:
            print(f"[OCR] PDF to Excel error: {e}")
            return False
    
    def full_convert(self, input_file, target_format='excel'):
        """Tổng hợp chuyển đổi"""
        if not self.ocr_available:
            print("[OCR] OCR not available")
            return None
        
        import pandas as pd
        
        ext = os.path.splitext(input_file)[1].lower()
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        timestamp = pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')
        
        os.makedirs('static/exports', exist_ok=True)
        
        if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            if target_format == 'excel':
                output_path = f"static/exports/{base_name}_{timestamp}.xlsx"
                success = self.process_image_to_excel(input_file, output_path)
                return output_path if success else None
            else:
                output_path = f"static/exports/{base_name}_{timestamp}.docx"
                text = self.process_image_to_text(input_file)
                if Document and text:
                    doc = Document()
                    doc.add_heading('OCR - PC06', 0)
                    for para in text.split('\n'):
                        if para.strip():
                            doc.add_paragraph(para)
                    doc.save(output_path)
                    return output_path
                return None
        
        elif ext == '.pdf':
            if target_format == 'excel':
                output_path = f"static/exports/{base_name}_{timestamp}.xlsx"
                success = self.process_pdf_to_excel(input_file, output_path)
                return output_path if success else None
            else:
                output_path = f"static/exports/{base_name}_{timestamp}.docx"
                success = self.process_pdf_to_word(input_file, output_path)
                return output_path if success else None
        
        return None

# Singleton
ocr_system = PC06_Tesseract_OCR()
