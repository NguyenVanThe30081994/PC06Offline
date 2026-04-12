import os
import io
import json
import base64
from flask import Blueprint, request, jsonify, send_file, render_template
from werkzeug.utils import secure_filename
import google.generativeai as genai

# Cấu hình Gemini API
GEMINI_API_KEY = "AIzaSyDVOb30nKEAJMNHH6pFX2xUdRULrBtE7C4"
genai.configure(api_key=GEMINI_API_KEY)

# Sử dụng model Gemini Pro Vision cho OCR và phân tích
model = genai.GenerativeModel('gemini-1.5-pro')

convert_bp = Blueprint('convert', __name__)

# Configuration
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'uploads')
CONVERT_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'tmp', 'convert')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERT_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@convert_bp.route('/convert')
def index():
    """Render convert page"""
    return render_template('convert.html')

@convert_bp.route('/convert/process', methods=['POST'])
def process():
    """Handle file conversion using Gemini AI"""
    if 'file' not in request.files:
        return jsonify({'error': 'Không tìm thấy file tải lên'}), 400
    
    file = request.files['file']
    convert_type = request.form.get('type', '')
    
    if file.filename == '':
        return jsonify({'error': 'Chưa chọn file'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Định dạng file không được hỗ trợ. Chỉ chấp nhận: JPG, PNG, PDF'}), 400
    
    filename = secure_filename(file.filename)
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)
    
    try:
        if convert_type == 'img_excel':
            return convert_image_to_excel_ai(filepath)
        elif convert_type == 'pdf_excel':
            return convert_pdf_to_excel_ai(filepath)
        elif convert_type == 'pdf_word':
            return convert_pdf_to_word_ai(filepath)
        elif convert_type == 'img_word':
            return convert_image_to_word_ai(filepath)
        elif convert_type == 'img_pdf':
            return convert_image_to_pdf_simple(filepath)
        else:
            return jsonify({'error': 'Loại chuyển đổi không hợp lệ'}), 400
    except Exception as e:
        print(f"Conversion error: {e}")
        return jsonify({'error': f'Lỗi xử lý: {str(e)}'}), 500
    finally:
        # Clean up uploaded file
        if os.path.exists(filepath):
            try:
                os.remove(filepath)
            except:
                pass

def convert_image_to_excel_ai(filepath):
    """Convert image to Excel using Gemini AI"""
    try:
        from openpyxl import Workbook
        
        # Đọc và mã hóa ảnh
        with open(filepath, 'rb') as f:
            image_data = f.read()
        
        # Gửi prompt cho Gemini phân tích bảng
        prompt = """Phân tích hình ảnh này và trích xuất dữ liệu bảng.
Trả về kết quả dưới dạng JSON với format:
```json
{"rows": [["cột 1", "cột 2", ...], ["dòng 1", "dòng 1", ...], ...]}
```
Chỉ trả về JSON, không giải thích thêm."""
        
        # Gọi Gemini với ảnh
        response = model.generate_content([
            prompt,
            {"mime_type": "image/jpeg", "data": image_data}
        ])
        
        # Parse kết quả
        text = response.text
        # Tìm và parse JSON
        start = text.find('{')
        end = text.rfind('}') + 1
        if start >= 0 and end > start:
            json_str = text[start:end]
            data = json.loads(json_str)
        else:
            return jsonify({'error': 'Không phân tích được dữ liệu từ ảnh'}), 400
        
        # Tạo Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "Bảng dữ liệu"
        
        rows = data.get('rows', [])
        for row_idx, row_data in enumerate(rows, 1):
            for col_idx, cell_data in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=cell_data)
        
        # Lưu vào buffer
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='bang_du_lieu.xlsx'
        )
        
    except Exception as e:
        print(f"Image to Excel AI error: {e}")
        return jsonify({'error': f'Lỗi chuyển đổi: {str(e)}'}), 500

def convert_pdf_to_excel_ai(filepath):
    """Convert PDF to Excel using Gemini AI"""
    try:
        from openpyxl import Workbook
        
        # Đọc file PDF
        with open(filepath, 'rb') as f:
            pdf_data = f.read()
        
        # Gửi prompt cho Gemini
        prompt = """Phân tích file PDF này và trích xuất tất cả các bảng dữ liệu.
Trả về kết quả dưới dạng JSON với format:
```json
{"tables": [[["cột 1", "cột 2"], ["dòng 1", "dòng 1"]], ...]}
```
Chỉ trả về JSON."""
        
        # Gọi Gemini với PDF
        response = model.generate_content([
            prompt,
            {"mime_type": "application/pdf", "data": pdf_data}
        ])
        
        # Parse kết quả
        text = response.text
        start = text.find('{')
        end = text.rfind('}') + 1
        if start >= 0 and end > start:
            json_str = text[start:end]
            data = json.loads(json_str)
        else:
            return jsonify({'error': 'Không phân tích được dữ liệu từ PDF'}), 400
        
        # Tạo Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "Bảng dữ liệu"
        
        tables = data.get('tables', [])
        row_count = 1
        for table in tables:
            for row in table:
                for col_idx, cell in enumerate(row, 1):
                    ws.cell(row=row_count, column=col_idx, value=cell if cell else "")
                row_count += 1
            row_count += 1
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='bang_tu_pdf.xlsx'
        )
        
    except Exception as e:
        print(f"PDF to Excel AI error: {e}")
        return jsonify({'error': f'Lỗi chuyển đổi: {str(e)}'}), 500

def convert_pdf_to_word_ai(filepath):
    """Convert PDF to Word using Gemini AI"""
    try:
        from docx import Document
        
        # Đọc file PDF
        with open(filepath, 'rb') as f:
            pdf_data = f.read()
        
        # Gửi prompt cho Gemini
        prompt = """Phân tích file PDF này và trích xuất toàn bộ nội dung văn bản.
Giữ nguyên định dạng và cấu trúc tài liệu.
Trả về nội dung dạng markdown."""
        
        # Gọi Gemini
        response = model.generate_content([
            prompt,
            {"mime_type": "application/pdf", "data": pdf_data}
        ])
        
        # Tạo Word document
        doc = Document()
        doc.add_heading('Nội dung tài liệu', 0)
        
        # Parse markdown và thêm vào Word
        lines = response.text.split('\n')
        for line in lines:
            if line.strip():
                # Kiểm tra nếu là heading
                if line.startswith('# '):
                    doc.add_heading(line[2:], 1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], 2)
                else:
                    doc.add_paragraph(line)
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='tai_lieu.docx'
        )
        
    except Exception as e:
        print(f"PDF to Word AI error: {e}")
        return jsonify({'error': f'Lỗi chuyển đổi: {str(e)}'}), 500

def convert_image_to_word_ai(filepath):
    """OCR image to Word using Gemini AI"""
    try:
        from docx import Document
        
        # Đọc và mã hóa ảnh
        with open(filepath, 'rb') as f:
            image_data = f.read()
        
        # Gửi prompt cho Gemini OCR
        prompt = """Nhận dạng toàn bộ văn bản trong hình ảnh này.
Trả về nội dung văn bản giữ nguyên format và xuống dòng.
Hỗ trợ tiếng Việt có dấu."""
        
        # Gọi Gemini
        response = model.generate_content([
            prompt,
            {"mime_type": "image/jpeg", "data": image_data}
        ])
        
        # Tạo Word document
        doc = Document()
        doc.add_heading('Văn bản nhận dạng từ ảnh', 0)
        
        # Thêm nội dung
        paragraphs = response.text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='van_ban.docx'
        )
        
    except Exception as e:
        print(f"Image to Word AI error: {e}")
        return jsonify({'error': f'Lỗi chuyển đổi: {str(e)}'}), 500

def convert_image_to_pdf_simple(filepath):
    """Convert image to PDF using PIL (lightweight)"""
    try:
        from PIL import Image
        
        # Mở ảnh
        img = Image.open(filepath)
        
        # Chuyển đổi sang RGB nếu cần
        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            if img.mode in ('RGBA', 'LA'):
                background.paste(img, mask=img.split()[-1])
                img = background
            else:
                img = img.convert('RGB')
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Lưu ra PDF
        output = io.BytesIO()
        img.save(output, format='PDF')
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='chuyen_doi.pdf'
        )
        
    except Exception as e:
        print(f"Image to PDF error: {e}")
        return jsonify({'error': f'Lỗi chuyển đổi ảnh sang PDF: {str(e)}'}), 500
